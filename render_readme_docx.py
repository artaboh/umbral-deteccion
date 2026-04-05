from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


@dataclass
class Block:
  kind: str
  level: int = 0
  text: str = ""
  items: list[str] | None = None
  lines: list[str] | None = None
  lang: str = ""


def parse_markdown(md: str) -> list[Block]:
  lines = md.splitlines()
  blocks: list[Block] = []
  i = 0

  while i < len(lines):
    raw = lines[i].rstrip("\n")
    line = raw.strip()

    if not line:
      i += 1
      continue

    if line.startswith("```"):
      lang = line[3:].strip()
      i += 1
      code_lines: list[str] = []
      while i < len(lines) and not lines[i].strip().startswith("```"):
        code_lines.append(lines[i].rstrip("\n"))
        i += 1
      if i < len(lines):
        i += 1
      blocks.append(Block(kind="code", lines=code_lines, lang=lang))
      continue

    m_h = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m_h:
      blocks.append(Block(kind="heading", level=len(m_h.group(1)), text=m_h.group(2).strip()))
      i += 1
      continue

    if re.match(r"^[-*+]\s+", line):
      items: list[str] = []
      while i < len(lines):
        li = lines[i].strip()
        if not re.match(r"^[-*+]\s+", li):
          break
        items.append(re.sub(r"^[-*+]\s+", "", li))
        i += 1
      blocks.append(Block(kind="ulist", items=items))
      continue

    if re.match(r"^\d+\.\s+", line):
      items: list[str] = []
      while i < len(lines):
        li = lines[i].strip()
        if not re.match(r"^\d+\.\s+", li):
          break
        items.append(re.sub(r"^\d+\.\s+", "", li))
        i += 1
      blocks.append(Block(kind="olist", items=items))
      continue

    para = [line]
    i += 1
    while i < len(lines):
      nxt = lines[i].strip()
      if (
        not nxt
        or nxt.startswith("```")
        or re.match(r"^(#{1,6})\s+", nxt)
        or re.match(r"^[-*+]\s+", nxt)
        or re.match(r"^\d+\.\s+", nxt)
      ):
        break
      para.append(nxt)
      i += 1
    blocks.append(Block(kind="paragraph", text=" ".join(para)))

  return blocks


def ensure_code_style(doc: Document) -> None:
  styles = doc.styles
  if "CodeBlock" in styles:
    return
  style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
  style.base_style = styles["Normal"]
  font = style.font
  font.name = "Consolas"
  font.size = Pt(9.5)
  font.color.rgb = RGBColor(30, 30, 30)
  p = style.paragraph_format
  p.space_before = Pt(4)
  p.space_after = Pt(4)
  p.left_indent = Pt(14)
  p.right_indent = Pt(8)
  p.line_spacing = 1.05


def set_code_background(paragraph) -> None:
  p_pr = paragraph._p.get_or_add_pPr()
  shd = OxmlElement("w:shd")
  shd.set(qn("w:val"), "clear")
  shd.set(qn("w:color"), "auto")
  shd.set(qn("w:fill"), "F2F4F7")
  p_pr.append(shd)


def add_inline_markdown(paragraph, text: str) -> None:
  parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
  for part in parts:
    if not part:
      continue
    if part.startswith("`") and part.endswith("`"):
      run = paragraph.add_run(part[1:-1])
      run.font.name = "Consolas"
      run.font.size = Pt(10)
      run.font.color.rgb = RGBColor(60, 60, 60)
    elif part.startswith("**") and part.endswith("**"):
      run = paragraph.add_run(part[2:-2])
      run.bold = True
    else:
      paragraph.add_run(part)


def render_docx(md_path: Path, docx_path: Path) -> None:
  md = md_path.read_text(encoding="utf-8")
  blocks = parse_markdown(md)
  doc = Document()

  normal = doc.styles["Normal"]
  normal.font.name = "Calibri"
  normal.font.size = Pt(11)
  normal.paragraph_format.space_after = Pt(6)

  ensure_code_style(doc)

  for block in blocks:
    if block.kind == "heading":
      level = min(max(block.level, 1), 3)
      p = doc.add_paragraph(style=f"Heading {level}")
      p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
      p.paragraph_format.space_after = Pt(4)
      add_inline_markdown(p, block.text)
      continue

    if block.kind == "paragraph":
      p = doc.add_paragraph(style="Normal")
      add_inline_markdown(p, block.text)
      continue

    if block.kind == "ulist":
      for item in block.items or []:
        p = doc.add_paragraph(style="List Bullet")
        add_inline_markdown(p, item)
      continue

    if block.kind == "olist":
      for item in block.items or []:
        p = doc.add_paragraph(style="List Number")
        add_inline_markdown(p, item)
      continue

    if block.kind == "code":
      if block.lang:
        tag = doc.add_paragraph(style="CodeBlock")
        run = tag.add_run(f"codigo: {block.lang}")
        run.bold = True
        run.font.color.rgb = RGBColor(75, 90, 120)
        set_code_background(tag)
      for line in (block.lines or [""]):
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line if line else " ")
        set_code_background(p)
      continue

  doc.save(docx_path)


def main() -> None:
  root = Path(__file__).resolve().parent
  md_path = root / "README.md"
  docx_path = root / "README.docx"
  render_docx(md_path, docx_path)
  print(f"DOCX generado: {docx_path}")


if __name__ == "__main__":
  main()
